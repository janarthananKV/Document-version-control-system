#!/usr/bin/env python3

import hashlib
import zipfile
from io import BytesIO
import tempfile
import subprocess
import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import xml.etree.ElementTree as ET
from PIL import Image
try:
    import fitz  # PyMuPDF
except Exception:  # optional runtime dependency
    fitz = None
import io
import difflib

def extract_image_metadata(docx_bytes):
    """
    Extract comprehensive image metadata including content hash, dimensions, 
    file info, and position within document.
    """
    images = []
    
    with zipfile.ZipFile(BytesIO(docx_bytes)) as z:
        # Get document.xml
        doc_xml = z.read("word/document.xml")
        root = ET.fromstring(doc_xml)
        
        # Find all image references
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
              "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
              "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
              "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
        
        # Find all embedded images
        for rel_id, img_data in _extract_embedded_images(z, root, ns):
            img_info = {
                'rel_id': rel_id,
                'content_hash': hashlib.md5(img_data).hexdigest(),
                'size_bytes': len(img_data),
                'position': _get_image_position(root, rel_id, ns),
                'dimensions': _get_image_dimensions(img_data),
                'format': _get_image_format(img_data)
            }
            images.append(img_info)
            
        # Also check for inline shapes (python-docx method)
        try:
            doc = Document(BytesIO(docx_bytes))
            for i, shape in enumerate(doc.inline_shapes):
                if hasattr(shape, 'image'):
                    # This is a more reliable way to get actual image data
                    img_info = {
                        'shape_index': i,
                        'width': shape.width,
                        'height': shape.height,
                        'type': str(shape.type),
                        'position': f"inline_shape_{i}",
                        'content_hash': _get_inline_shape_hash(shape),
                        'size_bytes': _get_inline_shape_size(shape)
                    }
                    images.append(img_info)
        except Exception as e:
            print(f"Warning: Could not extract inline shapes: {e}")
    
    # Try to enrich with exact page numbers via rendered PDF if available
    try:
        images = _enrich_positions_with_pdf_pages(images, docx_bytes)
    except Exception as e:
        # best-effort; ignore failures
        pass
    return images

def _extract_embedded_images(zip_file, doc_root, ns):
    """Extract embedded image data from DOCX relationships."""
    images = []
    
    # Read relationships
    try:
        rels_xml = zip_file.read("word/_rels/document.xml.rels")
        rels_root = ET.fromstring(rels_xml)
        
        for rel in rels_root.findall(".//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"):
            rel_type = rel.get("Type")
            if "image" in rel_type.lower():
                target = rel.get("Target")
                rel_id = rel.get("Id")
                
                try:
                    # Read the actual image file
                    img_path = f"word/{target}"
                    img_data = zip_file.read(img_path)
                    images.append((rel_id, img_data))
                except KeyError:
                    continue
    except KeyError:
        # No relationships file found
        pass
    
    return images

def _get_image_position(doc_root, rel_id, ns):
    """Return approximate position as page and paragraph: "page_X:paragraph_Y".

    We approximate pages by counting manual page breaks (w:br w:type="page").
    """
    r_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'

    # Build paragraph index -> page number mapping (best-effort)
    # We consider three signals that start a new page AFTER the paragraph:
    # 1) manual page break: w:br w:type="page"
    # 2) Word's last rendered page break: w:lastRenderedPageBreak
    # 3) section break that forces next page: w:sectPr/w:type[@w:val='nextPage']
    paragraphs = doc_root.findall('.//w:p', ns)
    page_by_para = {}
    current_page = 1
    w_type_attr = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type'
    w_val_attr = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'
    for idx, p in enumerate(paragraphs):
        page_by_para[idx] = current_page
        # 1) manual page break
        has_manual_break = any(br.get(w_type_attr) == 'page' for br in p.findall('.//w:br', ns))
        # 2) last rendered page break (layout-provided when Word saved the doc)
        has_rendered_break = p.find('.//w:lastRenderedPageBreak', ns) is not None
        # 3) section break with nextPage type
        has_nextpage_sect = False
        sect = p.find('.//w:sectPr', ns)
        if sect is not None:
            stype = sect.find('.//w:type', ns)
            if stype is not None and stype.get(w_val_attr) == 'nextPage':
                has_nextpage_sect = True
        if has_manual_break or has_rendered_break or has_nextpage_sect:
            current_page += 1

    # Find paragraph containing the image
    for idx, p in enumerate(paragraphs):
        for blip in p.findall('.//a:blip', ns):
            if blip.get(r_ns) == rel_id:
                page = page_by_para.get(idx, 1)
                return f"page_{page}:paragraph_{idx}"
    return "unknown_position"

def _get_image_dimensions(img_data):
    """Get image dimensions using PIL."""
    try:
        img = Image.open(io.BytesIO(img_data))
        return {"width": img.width, "height": img.height}
    except Exception:
        return {"width": 0, "height": 0}

def _get_image_format(img_data):
    """Detect image format."""
    try:
        img = Image.open(io.BytesIO(img_data))
        return img.format
    except Exception:
        return "unknown"

def _get_inline_shape_hash(shape):
    """Get hash for inline shape image."""
    try:
        if hasattr(shape, 'image'):
            return hashlib.md5(shape.image.blob).hexdigest()
    except Exception:
        pass
    return "no_hash"

def _get_inline_shape_size(shape):
    """Get size of inline shape image."""
    try:
        if hasattr(shape, 'image'):
            return len(shape.image.blob)
    except Exception:
        pass
    return 0

def _enrich_positions_with_pdf_pages(images, docx_bytes):
    """Best-effort: render DOCX to PDF and try to assign exact page numbers.

    Requirements at runtime: Windows with Word (for docx2pdf) or LibreOffice, and PyMuPDF.
    Implementation here uses PyMuPDF to scan rendered pages for embedded images and
    matches them using perceptual hash (pHash) against our image list.
    """
    if fitz is None:
        return images

    tmpdir = tempfile.mkdtemp(prefix="dvcs_imgpos_")
    docx_path = os.path.join(tmpdir, "doc.docx")
    pdf_path = os.path.join(tmpdir, "out.pdf")

    try:
        # Write DOCX
        with open(docx_path, 'wb') as f:
            f.write(docx_bytes)

        # Convert DOCX->PDF: try docx2pdf via Word (Windows) else soffice
        converted = False
        try:
            import docx2pdf  # type: ignore
            docx2pdf.convert(docx_path, pdf_path)
            converted = os.path.exists(pdf_path)
        except Exception:
            pass

        if not converted:
            import shutil
            soffice = shutil.which('soffice') or shutil.which('libreoffice')
            if soffice:
                subprocess.run([soffice, '--headless', '--convert-to', 'pdf', '--outdir', tmpdir, docx_path],
                               check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                # LibreOffice names output based on input name
                candidate = os.path.join(tmpdir, 'doc.pdf')
                pdf_path = candidate if os.path.exists(candidate) else pdf_path
                converted = os.path.exists(pdf_path)

        if not converted or not os.path.exists(pdf_path):
            return images

        # Scan rendered PDF pages for images and match by dimensions
        doc = fitz.open(pdf_path)
        try:
            for page_index in range(len(doc)):
                page = doc[page_index]
                img_list = page.get_images(full=True)
                for xref_info in img_list:
                    xref = xref_info[0]
                    pix = fitz.Pixmap(doc, xref)
                    try:
                        width, height = pix.width, pix.height
                    except Exception:
                        continue
                    candidates = []
                    for img in images:
                        dims = img.get('dimensions')
                        if dims and dims.get('width') == width and dims.get('height') == height:
                            candidates.append(img)
                    if len(candidates) == 1:
                        candidates[0]['position'] = f"page_{page_index+1}"
        finally:
            doc.close()
        return images
    except Exception:
        return images
    finally:
        try:
            if os.path.exists(docx_path):
                os.remove(docx_path)
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
            if os.path.exists(tmpdir):
                os.rmdir(tmpdir)
        except Exception:
            pass

def compare_images_detailed(current_images, other_images):
    """
    Compare two sets of images and return detailed differences.
    """
    changes = {
        'added': [],
        'removed': [],
        'modified': [],
        'unchanged': []
    }
    
    # Create lookup dictionaries
    current_dict = {img['content_hash']: img for img in current_images}
    other_dict = {img['content_hash']: img for img in other_images}
    
    # Find added images
    for hash_val, img in other_dict.items():
        if hash_val not in current_dict:
            changes['added'].append(img)
    
    # Find removed images
    for hash_val, img in current_dict.items():
        if hash_val not in other_dict:
            changes['removed'].append(img)
    
    # Find unchanged images
    for hash_val in current_dict:
        if hash_val in other_dict:
            changes['unchanged'].append(current_dict[hash_val])
    
    return changes

def describe_images_enhanced(docx_bytes):
    """
    Enhanced version of the original describe_images function.
    Returns a list of image descriptors for comparison.
    """
    images = extract_image_metadata(docx_bytes)
    descriptors = []
    
    for img in images:
        if 'dimensions' in img:
            desc = f"{img['dimensions']['width']}x{img['dimensions']['height']}_{img['format']}_{img['size_bytes']}b"
        else:
            desc = f"{img['width']}x{img['height']}_{img.get('type', 'unknown')}"
        descriptors.append(desc)
    
    return descriptors

def generate_enhanced_highlighted_copy(current_bytes, other_bytes, output_path="enhanced_highlighted_copy.docx"):
    """
    Generate a new DOCX with enhanced image difference highlighting.
    """
    doc_current = Document(BytesIO(current_bytes))
    doc_other = Document(BytesIO(other_bytes))
    
    # Extract detailed image information
    current_images = extract_image_metadata(current_bytes)
    other_images = extract_image_metadata(other_bytes)
    
    # Compare images
    image_changes = compare_images_detailed(current_images, other_images)
    
    # Create new document with enhanced reporting
    new_doc = Document()
    
    # Add title
    title = new_doc.add_heading('Document Version Comparison Report', 0)
    
    # Text differences (existing logic)
    current_lines = [p.text for p in doc_current.paragraphs]
    other_lines = [p.text for p in doc_other.paragraphs]
    diff = list(difflib.ndiff(current_lines, other_lines))
    
    new_doc.add_heading('Text Changes', level=1)
    
    # --- Paragraph/Text comparison ---
    i = 0
    while i < len(diff):
        line = diff[i]
        if line.startswith("- ") and i + 1 < len(diff) and diff[i + 1].startswith("+ "):
            p = new_doc.add_paragraph()
            run = p.add_run(diff[i + 1][2:])
            run.font.color.rgb = RGBColor(255, 0, 0)  # replaced → red
            i += 2
            continue
        elif line.startswith("+ "):  # added
            p = new_doc.add_paragraph()
            run = p.add_run(line[2:])
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        elif line.startswith("- "):  # deleted
            p = new_doc.add_paragraph()
            run = p.add_run("[Removed text]")
            run.font.color.rgb = RGBColor(128, 128, 128)
        elif not line.startswith("? "):
            new_doc.add_paragraph(line[2:])
        i += 1
    
    # Enhanced image reporting
    new_doc.add_heading('Image Changes Analysis', level=1)
    
    if image_changes['added']:
        new_doc.add_heading('Added Images', level=2)
        for img in image_changes['added']:
            p = new_doc.add_paragraph()
            if 'dimensions' in img:
                run = p.add_run(f"✓ Added: {img['format']} image ({img['dimensions']['width']}x{img['dimensions']['height']}) - {img['size_bytes']} bytes")
            else:
                run = p.add_run(f"✓ Added: Image ({img['width']}x{img['height']}) - {img['size_bytes']} bytes")
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    
    if image_changes['removed']:
        new_doc.add_heading('Removed Images', level=2)
        for img in image_changes['removed']:
            p = new_doc.add_paragraph()
            if 'dimensions' in img:
                run = p.add_run(f"✗ Removed: {img['format']} image ({img['dimensions']['width']}x{img['dimensions']['height']}) - {img['size_bytes']} bytes")
            else:
                run = p.add_run(f"✗ Removed: Image ({img['width']}x{img['height']}) - {img['size_bytes']} bytes")
            run.font.color.rgb = RGBColor(255, 0, 0)
    
    if image_changes['unchanged']:
        new_doc.add_heading('Unchanged Images', level=2)
        for img in image_changes['unchanged']:
            p = new_doc.add_paragraph()
            if 'dimensions' in img:
                run = p.add_run(f"• Unchanged: {img['format']} image ({img['dimensions']['width']}x{img['dimensions']['height']})")
            else:
                run = p.add_run(f"• Unchanged: Image ({img['width']}x{img['height']})")
            run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Summary
    summary_para = new_doc.add_paragraph()
    summary_text = f"Summary: {len(image_changes['added'])} added, {len(image_changes['removed'])} removed, {len(image_changes['unchanged'])} unchanged images"
    summary_para.add_run(summary_text)
    
    new_doc.save(output_path)
    print(f"[OK] Enhanced highlighted copy saved: {output_path}")
    print(f"Image changes: {len(image_changes['added'])} added, {len(image_changes['removed'])} removed, {len(image_changes['unchanged'])} unchanged")
    
    return image_changes
