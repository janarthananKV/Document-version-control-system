#!/usr/bin/env python3
import os
import json
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn
from deepdiff import DeepDiff
import google.generativeai as genai

from dvcs import DVCS  # import your DVCS class
from image_utils import extract_image_metadata, compare_images_detailed

# --- API key management helpers (avoid typing every time) ---
_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
_DEFAULT_KEY_FILE = os.path.join(_BASE_DIR, ".dvcs_gemini_key")

def resolve_gemini_key(provided_key: str | None) -> str | None:
    """Return an API key using precedence: provided > env > keyfile.

    - Env var name: GEMINI_API_KEY
    - Key file path: ~/.dvcs_gemini_key (plain text, single line)
    """
    if provided_key:
        return provided_key.strip()
    env_key = os.getenv("GEMINI_API_KEY")
    if env_key:
        return env_key.strip()
    try:
        if os.path.exists(_DEFAULT_KEY_FILE):
            with open(_DEFAULT_KEY_FILE, "r", encoding="utf-8") as f:
                data = f.read().strip()
                return data or None
    except Exception:
        pass

def save_gemini_key(key: str) -> None:
    """Persist API key to the default key file in the user's home directory."""
    try:
        os.makedirs(os.path.dirname(_DEFAULT_KEY_FILE), exist_ok=True)
    except Exception:
        # Parent may already exist or be the home itself
        pass
    with open(_DEFAULT_KEY_FILE, "w", encoding="utf-8") as f:
        f.write((key or "").strip())

def extract_pages_by_breaks(docx_bytes):
    """
    Splits a DOCX document into pages based on manual page breaks (Ctrl+Enter).
    Returns a dictionary: {page_number: {"paragraphs": [], "tables": [], "images": []}}
    """
    doc = Document(BytesIO(docx_bytes))
    pages = {}
    current_page = 1
    pages[current_page] = {"paragraphs": [], "tables": [], "images": []}

    for p in doc.paragraphs:
        runs = []
        for r in p.runs:
            runs.append({
                "text": r.text,
                "bold": r.bold,
                "italic": r.italic,
                "underline": r.underline
            })
            for br in r._element.findall(".//w:br", r._element.nsmap):
                if br.get(qn("w:type")) == "page":
                    current_page += 1
                    if current_page not in pages:
                        pages[current_page] = {"paragraphs": [], "tables": [], "images": []}
        pages[current_page]["paragraphs"].append(runs)

    for t in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in t.rows]
        pages[current_page]["tables"].append(rows)

    for shape in doc.inline_shapes:
        pages[current_page]["images"].append({
            "width": shape.width,
            "height": shape.height,
            "type": str(shape.type)
        })
    
    # Add enhanced image metadata
    try:
        enhanced_images = extract_image_metadata(docx_bytes)
        pages[current_page]["enhanced_images"] = enhanced_images
    except Exception as e:
        print(f"Warning: Could not extract enhanced image metadata: {e}")
        pages[current_page]["enhanced_images"] = []

    return pages


from docx import Document
from io import BytesIO
from docx.oxml.ns import qn
from deepdiff import DeepDiff

# ============================
# ULTRA-DETAILED DOCX EXTRACTOR
# ============================

def extract_runs(paragraph):
    runs = []
    for r in paragraph.runs:
        run_info = {
            "text": r.text,
            "bold": r.bold,
            "italic": r.italic,
            "underline": r.underline,
            "font_size": r.font.size.pt if r.font.size else None,
            "font_name": r.font.name,
            "color": r.font.color.rgb.__str__() if r.font.color and r.font.color.rgb else None,
            "highlight": str(r.font.highlight_color) if r.font.highlight_color else None,
        }
        runs.append(run_info)
    return runs


def extract_paragraph_info(paragraph):
    alignment = None
    if paragraph.alignment is not None:
        alignment = paragraph.alignment.real

    return {
        "style": paragraph.style.name if paragraph.style else None,
        "alignment": alignment,
        "runs": extract_runs(paragraph)
    }


def extract_table_info(table):
    table_data = []
    for row in table.rows:
        row_info = []
        for cell in row.cells:
            row_info.append({
                "text": cell.text,
                "paragraphs": [extract_paragraph_info(p) for p in cell.paragraphs]
            })
        table_data.append(row_info)
    return table_data


def extract_doc_structure(docx_bytes):
    doc = Document(BytesIO(docx_bytes))

    structure = {
        "paragraphs": [],
        "tables": [],
        "images": []
    }

    for p in doc.paragraphs:
        structure["paragraphs"].append(extract_paragraph_info(p))

    for t in doc.tables:
        structure["tables"].append(extract_table_info(t))

    for shape in doc.inline_shapes:
        structure["images"].append({
            "width": shape.width,
            "height": shape.height,
            "type": str(shape.type)
        })

    return structure


def diff_docx_bytes(v1_bytes, v2_bytes):
    doc1 = extract_doc_structure(v1_bytes)
    doc2 = extract_doc_structure(v2_bytes)
    return DeepDiff(doc1, doc2, ignore_order=True).to_dict()


def summarize_snapshot(doc_name: str, v1: int, v2: int,
                       use_llm=False, llm_api_key=None, output_dir=None):

    dvcs = DVCS(doc_name, "docx")

    # Reconstruct bytes
    b1 = dvcs._reconstruct_bytes(v1)
    b2 = dvcs._reconstruct_bytes(v2)

    # ===============================
    # 1. IMAGE ANALYSIS (your part)
    # ===============================
    try:
        images_v1 = extract_image_metadata(b1)
        images_v2 = extract_image_metadata(b2)
        image_changes = compare_images_detailed(images_v1, images_v2)

        summary_text = (
            f"\n=== Enhanced Image Analysis ===\n"
            f"Images in v{v1}: {len(images_v1)}\n"
            f"Images in v{v2}: {len(images_v2)}\n"
            f"Added: {len(image_changes['added'])}\n"
            f"Removed: {len(image_changes['removed'])}\n"
            f"Unchanged: {len(image_changes['unchanged'])}\n"
        )
    except Exception as e:
        summary_text = f"\n[Image Analysis Error: {e}]\n"
        image_changes = {'added': [], 'removed': [], 'unchanged': []}

    # ======================================
    # 2. FULL DOCX STRUCTURE DIFF (NEW PART)
    # ======================================
    full_diff = diff_docx_bytes(b1, b2)
    full_diff = json.loads(json.dumps(full_diff, default=str))


    if not full_diff:
        return summary_text + "\n=== No text/table/style changes detected ===\n"

    # ========================================================
    # 3. LLM Summarization (with deep diff + image context)
    # ========================================================
    if use_llm:
        effective_key = resolve_gemini_key(llm_api_key)
        if effective_key:
            try:
                genai.configure(api_key=effective_key)
            except Exception:
                pass

        image_context = ""
        if image_changes["added"] or image_changes["removed"]:
            image_context = (
                f"\nImage changes: {len(image_changes['added'])} added, "
                f"{len(image_changes['removed'])} removed."
            )

        prompt = (
            f"You are a professional document-change summarizer.\n\n"
            f"Summarize ALL detected changes:\n"
            f"- Paragraph text changes\n"
            f"- Run-level changes (bold, italic, underline, font size, font name, color)\n"
            f"- Paragraph alignment/format changes\n"
            f"- Table structure and cell-level changes\n"
            f"- Image changes\n\n"
            f"Write the summary in 6–10 bullet points.\n\n"
            f"Image info: {image_context}\n\n"
            f"Deep structured diff JSON:\n{json.dumps(full_diff, indent=2)}"
        )

        try:
            model = genai.GenerativeModel("gemini-2.5-flash")
            response = model.generate_content(prompt)
            summary_text += f"\n=== Document Summary ===\n{response.text}\n"
        except Exception as e:
            summary_text += f"\n[Gemini API ERROR: {e}]\nPrompt Preview:\n{prompt[:2000]}\n"
    else:
        summary_text += (
            "\n=== RAW STRUCTURED DIFF ===\n"
            + json.dumps(full_diff, indent=2)
            + "\n"
        )

    return summary_text




