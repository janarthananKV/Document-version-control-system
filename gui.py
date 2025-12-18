# import tkinter as tk
# from tkinter import filedialog, simpledialog, messagebox, ttk
# from pathlib import Path
# from dvcs import DVCS, RepoState, repo_dir_for
# from summarizer import summarize_snapshot


# class DVCSApp:
#     def __init__(self, root):
#         self.root = root
#         self.root.title("DOCX Version Control System")
#         self.dvcs = None
#         self.file_path = None

#         # Toolbar
#         toolbar = tk.Frame(root)
#         toolbar.pack(side=tk.TOP, fill=tk.X)

#         tk.Button(toolbar, text="Open File", command=self.open_file).pack(side=tk.LEFT, padx=5, pady=5)
#         tk.Button(toolbar, text="Commit", command=self.commit).pack(side=tk.LEFT, padx=5)
#         tk.Button(toolbar, text="Rollback", command=self.rollback).pack(side=tk.LEFT, padx=5)
#         tk.Button(toolbar, text="Summarize Changes", command=self.summarize_changes).pack(side=tk.LEFT, padx=5)

#         # History
#         self.tree = ttk.Treeview(
#             root, columns=("version", "kind", "date", "message"),
#             show="headings", selectmode="extended"
#         )
#         for col in self.tree["columns"]:
#             self.tree.heading(col, text=col.capitalize())
#         self.tree.pack(fill=tk.BOTH, expand=True)

#     def open_file(self):
#         file = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
#         if not file:
#             return
#         self.file_path = Path(file)

#         repo_dir = repo_dir_for(self.file_path)
#         state = RepoState.load(repo_dir)

#         if state is None:
#             if messagebox.askyesno("Init Repo", "No repo found. Initialize?"):
#                 msg = simpledialog.askstring("Commit Message", "Initial commit message:")
#                 self.dvcs = DVCS(file, "docx")
#                 self.dvcs.init(msg or "Initial version")
#             else:
#                 return
#         else:
#             self.dvcs = DVCS(file, state.document_type)

#         self.refresh_history()

#     def commit(self):
#         if not self.dvcs:
#             return
#         msg = simpledialog.askstring("Commit", "Enter commit message:")
#         self.dvcs.add(msg or "Update")
#         self.refresh_history()

#     def rollback(self):
#         if not self.dvcs:
#             return
#         item = self.tree.selection()
#         if not item:
#             messagebox.showerror("Error", "Select a version to rollback.")
#             return
#         version = int(self.tree.item(item, "values")[0])
#         if messagebox.askyesno("Confirm", f"Revert file to version {version}?"):
#             self.dvcs.revert(version)
#             messagebox.showinfo("Done", f"File reverted to v{version}.")

#     def summarize_changes(self):
#         if not self.dvcs:
#             return
#         selected = self.tree.selection()
#         if len(selected) != 2:
#             messagebox.showerror("Error", "Select exactly 2 versions.")
#             return

#         v1, v2 = [int(self.tree.item(i, "values")[0]) for i in selected]

#         # Ask if LLM summary is desired
#         use_llm = messagebox.askyesno("Use Gemini?", "Do you want a natural-language summary with Gemini?")
#         llm_key = None
#         if use_llm:
#             llm_key = simpledialog.askstring("Gemini API Key", "Enter your Gemini API key (or leave blank if configured globally):")

#         result = summarize_snapshot(str(self.file_path), v1, v2, use_llm=use_llm, llm_api_key=llm_key)
#         self.show_text_window(f"Summary v{v1} → v{v2}", result)


#     def show_text_window(self, title, text):
#         win = tk.Toplevel(self.root)
#         win.title(title)
#         txt = tk.Text(win, wrap="word")
#         txt.insert("1.0", text)
#         txt.pack(fill=tk.BOTH, expand=True)
#         scroll = tk.Scrollbar(win, command=txt.yview)
#         scroll.pack(side=tk.RIGHT, fill=tk.Y)
#         txt.config(yscrollcommand=scroll.set)

#     def refresh_history(self):
#         if not self.dvcs:
#             return
#         for row in self.tree.get_children():
#             self.tree.delete(row)
#         for v in self.dvcs.state.versions:
#             self.tree.insert("", "end", values=(v.version, v.kind, v.created_at, v.message))


# if __name__ == "__main__":
#     root = tk.Tk()
#     app = DVCSApp(root)
#     root.mainloop()


# import tkinter as tk
# from tkinter import ttk, filedialog, messagebox
# from tkinter.scrolledtext import ScrolledText
# from dvcs import DVCS
# from summarizer import summarize_snapshot

# class DVCS_GUI(tk.Tk):
#     def __init__(self):
#         super().__init__()
#         self.title("Docx Version Control System")
#         self.geometry("900x650")
#         self.configure(bg="#f5f5f5")

#         self.dvcs = None
#         self.file_path = None

#         style = ttk.Style()
#         style.configure("TButton", padding=6, font=("Helvetica", 10))
#         style.configure("TLabel", font=("Helvetica", 10))
#         style.configure("TFrame", background="#f5f5f5")
#         style.configure("TCombobox", padding=3)

#         self.create_widgets()

#     def create_widgets(self):
#         # --- File selection frame ---
#         file_frame = ttk.Frame(self, padding=10)
#         file_frame.pack(fill="x", padx=10, pady=5)

#         ttk.Label(file_frame, text="Document:").pack(side="left")
#         self.file_label = ttk.Label(file_frame, text="No file selected", width=50)
#         self.file_label.pack(side="left", padx=5)
#         ttk.Button(file_frame, text="Select File", command=self.select_file).pack(side="left", padx=5)

#         # --- Version selection frame ---
#         version_frame = ttk.Frame(self, padding=10, relief="groove")
#         version_frame.pack(fill="x", padx=10, pady=5)

#         ttk.Label(version_frame, text="Version 1:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
#         ttk.Label(version_frame, text="Version 2:").grid(row=0, column=2, padx=5, pady=5, sticky="e")

#         self.version1_var = tk.StringVar()
#         self.version2_var = tk.StringVar()
#         self.version1_cb = ttk.Combobox(version_frame, textvariable=self.version1_var, width=10, state="readonly")
#         self.version2_cb = ttk.Combobox(version_frame, textvariable=self.version2_var, width=10, state="readonly")
#         self.version1_cb.grid(row=0, column=1, padx=5, pady=5)
#         self.version2_cb.grid(row=0, column=3, padx=5, pady=5)

#         ttk.Button(version_frame, text="Rollback", command=self.rollback).grid(row=0, column=4, padx=5)
#         ttk.Button(version_frame, text="Summarize Changes", command=self.summarize_changes).grid(row=0, column=5, padx=5)

#         # LLM option
#         self.llm_var = tk.BooleanVar()
#         ttk.Checkbutton(version_frame, text="Use Gemini LLM", variable=self.llm_var).grid(row=1, column=0, columnspan=2, pady=5)
#         self.llm_key_var = tk.StringVar()
#         ttk.Entry(version_frame, textvariable=self.llm_key_var, width=40).grid(row=1, column=2, columnspan=3, padx=5, pady=5)

#         # --- Output text frame ---
#         output_frame = ttk.Frame(self, padding=10)
#         output_frame.pack(fill="both", expand=True, padx=10, pady=5)

#         self.output_text = ScrolledText(output_frame, wrap="word", font=("Consolas", 10))
#         self.output_text.pack(fill="both", expand=True)

#     # --- File selection ---
#     def select_file(self):
#         path = filedialog.askopenfilename(filetypes=[("DOCX files", "*.docx")])
#         if path:
#             self.file_path = path
#             self.file_label.config(text=path)
#             self.dvcs = DVCS(path, "docx")
#             self.populate_versions()

#     def populate_versions(self):
#         if self.dvcs:
#             versions = [str(v.version) for v in self.dvcs.state.versions]
#             self.version1_cb['values'] = versions
#             self.version2_cb['values'] = versions
#             if versions:
#                 self.version1_cb.current(0)
#                 self.version2_cb.current(len(versions)-1)

#     # --- Rollback ---
#     def rollback(self):
#         version = self.version1_var.get()
#         if not version:
#             messagebox.showerror("Error", "Select version to rollback")
#             return
#         try:
#             self.dvcs.revert(int(version))
#             messagebox.showinfo("Rollback", f"Reverted {self.file_path} to v{version}")
#         except Exception as e:
#             messagebox.showerror("Error", str(e))

#     # --- Summarize changes ---
#     def summarize_changes(self):
#         v1 = self.version1_var.get()
#         v2 = self.version2_var.get()
#         if not v1 or not v2:
#             messagebox.showerror("Error", "Select both versions")
#             return
#         use_llm = self.llm_var.get()
#         llm_key = self.llm_key_var.get() if use_llm else None

#         try:
#             summary = summarize_snapshot(self.file_path, int(v1), int(v2), use_llm=use_llm, llm_api_key=llm_key)
#             self.show_text_window(f"Summary v{v1} → v{v2}", summary)
#         except Exception as e:
#             messagebox.showerror("Error", str(e))

#     def show_text_window(self, title, text):
#         self.output_text.delete("1.0", tk.END)
#         self.output_text.insert(tk.END, text)

# if __name__ == "__main__":
#     app = DVCS_GUI()
#     app.mainloop()


import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
from dvcs import DVCS, repo_dir_for


from docx import Document
import difflib
import tkinter as tk

from docx import Document
import difflib

def extract_docx_elements(path):
    """Extract paragraphs, tables, and image markers from a DOCX file."""
    doc = Document(path)
    elements = []
    for block in doc.element.body:
        tag = block.tag.lower()
        if "p" in tag:
            # Paragraph
            text = "".join([t.text for t in block.xpath(".//w:t")])
            if text.strip():
                elements.append(f"TEXT: {text}")
        elif "tbl" in tag:
            # Table
            elements.append("[TABLE]")
        elif "drawing" in tag or "pic" in tag:
            elements.append("[IMAGE]")
    return elements


def diff_texts_with_media(path1, path2):
    """Generate diff including image and table placeholders."""
    elements1 = extract_docx_elements(path1)
    elements2 = extract_docx_elements(path2)
    diff = difflib.ndiff(elements1, elements2)
    return list(diff)


def get_text_from_docx(path):
    doc = Document(path)
    return [p.text for p in doc.paragraphs if p.text.strip()]

def show_diff_in_gui(root, diff):
    text_widget = tk.Text(root, wrap="word", font=("Poppins", 11))
    text_widget.pack(fill="both", expand=True)

    # Define tags for color
    text_widget.tag_config("added", foreground="green")
    text_widget.tag_config("removed", foreground="red")
    text_widget.tag_config("unchanged", foreground="black")
    text_widget.tag_config("table", foreground="orange", font=("Poppins", 11, "italic"))
    text_widget.tag_config("image", foreground="blue", font=("Poppins", 11, "italic"))

    for line in diff:
        content = line[2:]
        if line.startswith("+ "):
            if "[IMAGE]" in content:
                tag = "image"
            elif "[TABLE]" in content:
                tag = "table"
            else:
                tag = "added"
            text_widget.insert("end", content + "\n", tag)
        elif line.startswith("- "):
            if "[IMAGE]" in content:
                tag = "image"
            elif "[TABLE]" in content:
                tag = "table"
            else:
                tag = "removed"
            text_widget.insert("end", content + "\n", tag)
        else:
            text_widget.insert("end", content + "\n", "unchanged")



class DVCS_UI:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Version Control System")
        self.root.geometry("950x650")
        self.file_path = None
        self.dvcs = None

        self.create_widgets()

    def create_widgets(self):
        # Top frame: select file
        file_frame = ttk.Frame(self.root, padding=10)
        file_frame.pack(fill="x")

        ttk.Label(file_frame, text="Selected File:").pack(side="left")
        self.file_label = ttk.Label(file_frame, text="None", width=60, relief="sunken")
        self.file_label.pack(side="left", padx=5)

        ttk.Button(file_frame, text="Select DOCX", command=self.select_file).pack(side="left", padx=5)
        ttk.Button(file_frame, text="Init Repository", command=self.init_repo).pack(side="left", padx=5)
        ttk.Button(file_frame, text="Show History", command=self.show_history).pack(side="left", padx=5)

        # Version frame: rollback / commit / summarize
        version_frame = ttk.LabelFrame(self.root, text="Version Actions", padding=10)
        version_frame.pack(fill="x", padx=10, pady=10)

        # Rollback
        ttk.Label(version_frame, text="Rollback to Version:").grid(row=0, column=0, padx=5)
        self.version_var = tk.IntVar()
        self.version_combo = ttk.Combobox(version_frame, textvariable=self.version_var, width=10)
        self.version_combo.grid(row=0, column=1, padx=5)
        ttk.Button(version_frame, text="Rollback", command=self.rollback).grid(row=0, column=2, padx=5)

        # Commit
        ttk.Button(version_frame, text="Commit Changes", command=self.commit_changes).grid(row=0, column=3, padx=5)

        # Summarize
        ttk.Label(version_frame, text="Compare Versions:").grid(row=1, column=0, padx=5)
        self.v1_var = tk.IntVar()
        self.v2_var = tk.IntVar()
        self.v1_combo = ttk.Combobox(version_frame, textvariable=self.v1_var, width=10)
        self.v2_combo = ttk.Combobox(version_frame, textvariable=self.v2_var, width=10)
        self.v1_combo.grid(row=1, column=1, padx=5)
        self.v2_combo.grid(row=1, column=2, padx=5)
        ttk.Button(version_frame, text="Summarize Changes", command=self.summarize_changes).grid(row=1, column=3, padx=5)

        # Text area for output / summaries
        text_frame = ttk.LabelFrame(self.root, text="Output / Summary", padding=10)
        text_frame.pack(fill="both", expand=True, padx=10, pady=10)

        self.text_area = tk.Text(text_frame, wrap="word")
        self.text_area.pack(fill="both", expand=True)

        ttk.Button(version_frame, text="Compare Versions", command=self.compare_versions).grid(row=2, column=0, columnspan=2, pady=5)


    def select_file(self):
        path = filedialog.askopenfilename(filetypes=[("DOCX Files", "*.docx")])
        if path:
            self.file_path = path
            self.file_label.config(text=self.file_path)
            rdir = repo_dir_for(self.file_path)
            if rdir.exists():
                self.dvcs = DVCS(self.file_path, "docx")
            else:
                self.dvcs = None
            self.populate_versions()

    def init_repo(self):
        if not self.file_path:
            messagebox.showerror("Error", "Select a file first")
            return
        try:
            self.dvcs = DVCS(self.file_path, "docx")
            self.dvcs.init("Initial commit")
            messagebox.showinfo("Success", "Repository initialized!")
            self.populate_versions()
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def populate_versions(self):
        if not self.dvcs:
            return
        versions = [v.version for v in self.dvcs.state.versions]
        self.version_combo['values'] = versions
        self.v1_combo['values'] = versions
        self.v2_combo['values'] = versions

    def commit_changes(self):
        if not self.file_path or not self.dvcs:
            messagebox.showerror("Error", "No document selected or repo not initialized")
            return
        message = simpledialog.askstring("Commit Message", "Enter commit message:")
        if not message:
            return
        try:
            self.dvcs.add(message)
            messagebox.showinfo("Committed", f"Changes committed with message:\n{message}")
            self.populate_versions()
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def rollback(self):
        if not self.dvcs:
            messagebox.showerror("Error", "No document selected or repo not initialized")
            return
        version = self.version_var.get()
        if not version:
            messagebox.showerror("Error", "Select a version to rollback")
            return
        try:
            self.dvcs.revert(version)
            messagebox.showinfo("Reverted", f"Reverted to version {version}")
            self.populate_versions()
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def show_history(self):
        if not self.dvcs:
            messagebox.showerror("Error", "No document selected or repo not initialized")
            return
        history_window = tk.Toplevel(self.root)
        history_window.title(f"History - {self.file_path}")
        history_window.geometry("700x400")

        tree = ttk.Treeview(history_window, columns=("version", "kind", "message", "created_at"), show="headings")
        tree.heading("version", text="Version")
        tree.heading("kind", text="Kind")
        tree.heading("message", text="Message")
        tree.heading("created_at", text="Created At")

        for v in self.dvcs.state.versions:
            tree.insert("", "end", values=(v.version, v.kind, v.message, v.created_at))

        tree.pack(fill="both", expand=True)

    def summarize_changes(self):
        if not self.dvcs:
            messagebox.showerror("Error", "No document selected or repo not initialized")
            return
        v1 = self.v1_var.get()
        v2 = self.v2_var.get()
        if not v1 or not v2:
            messagebox.showerror("Error", "Select both versions to compare")
            return

        try:
            from summarizer import summarize_snapshot

            use_llm = messagebox.askyesno("Use Gemini LLM?", "Do you want to use Gemini for summary?")
            

            # Capture summary text instead of printing
            summary_text = summarize_snapshot(str(self.file_path), v1, v2, use_llm=use_llm, llm_api_key=None, output_dir=None)

            # Display in Text widget
            self.text_area.delete("1.0", tk.END)
            if summary_text:
                self.text_area.insert("1.0", summary_text)
            else:
                self.text_area.insert("1.0", f"No summary generated for v{v1} -> v{v2}")

        except Exception as e:
            messagebox.showerror("Error", str(e))
    

    def compare_versions(self):
        if not self.dvcs:
            messagebox.showerror("Error", "No document selected or repo not initialized")
            return

        v1 = self.v1_var.get()
        v2 = self.v2_var.get()
        if not v1 or not v2:
            messagebox.showerror("Error", "Select both versions to compare")
            return

        try:
            path_v1 = self.dvcs.get_version_file(v1)
            path_v2 = self.dvcs.get_version_file(v2)

            diff = diff_texts_with_media(path_v1, path_v2)

            diff_window = tk.Toplevel(self.root)
            diff_window.title(f"Diff: v{v1} vs v{v2}")
            show_diff_in_gui(diff_window, diff)

        except Exception as e:
            messagebox.showerror("Error", str(e))




if __name__ == "__main__":
    root = tk.Tk()
    app = DVCS_UI(root)
    root.mainloop()
